# import docx
# doc= docx.Document('WithOutImageQuestion.docx')
# fulltext=[]
# for para in doc.paragraphs:
#     fulltext.append(para.text)
# text = [i for i in fulltext if i.strip()]
# newtext=[]
# for val in text:
#     if "\t" in val:
#         x=val.replace("\t","")
#         newtext.append(x)
#     else:
#         newtext.append(val)
# print(newtext)
# print(newtext[:4])


# import xlsxwriter
# data=[
#     {
#     "name":"Steven",
#     "phone":"12345",
#     "email":"steven@gmail.com",
#     "country":"australia",
# },
#     {
#     "name":"philps",
#     "phone":"67089",
#     "email":"philphs@gmail.com",
#     "country":"canada",
# },
#     {
#     "name":"andrey",
#     "phone":"98076",
#     "email":"andrey@gmail.com",
#     "country":"california",
# }
# ]

# workbook = xlsxwriter.Workbook("AllProjectDetail.xlsx")
# worksheet = workbook.add_worksheet('firstsheet')

# worksheet.write(0, 0 ,"#")
# worksheet.write(0, 1 ,"name")
# worksheet.write(0, 2 ,"phone")
# worksheet.write(0, 3 ,"email")
# worksheet.write(0, 4 ,"country")

# for index, entry in enumerate(data):
#     worksheet.write(index+1, 0, str(index))
#     worksheet.write(index+1, 1, entry['name'])
#     worksheet.write(index+1, 2, entry['phone'])
#     worksheet.write(index+1, 3, entry['email'])
#     worksheet.write(index+1, 4, entry['country'])
    
# workbook.close()

# ***************************************************************************

import docx
doc= docx.Document('WithOutImageQuestion.docx')
fulltext=[]
for para in doc.paragraphs:
    fulltext.append(para.text)
text = [i for i in fulltext if i.strip()]
newtext=[]
for val in text:
    if "\t" in val:
        x=val.replace("\t","")
        newtext.append(x)
    else:
        newtext.append(val)
# print(newtext)
print(newtext[:4])



